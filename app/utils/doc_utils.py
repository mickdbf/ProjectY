import os
import re
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def create_doc_from_text(text: str, title: str = "AI_Generated_Notes") -> str:
    """
    Converts EasyNote Markup (ENM) structured text into a beautifully formatted Word document.
    Handles #SECTION, #SUBSECTION, #POINT, #NOTE, #CODE, #QUOTE, and #DIVIDER tags.
    """

    # -------------------------------
    # Setup document
    # -------------------------------
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    base_font = "Calibri"
    heading_color = RGBColor(41, 73, 140)
    body_color = RGBColor(45, 45, 45)
    soft_gray = RGBColor(120, 120, 120)

    # -------------------------------
    # Title + header
    # -------------------------------
    title_para = doc.add_paragraph(title)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title_para.runs:
        run.font.name = base_font
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = heading_color

    timestamp = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    )
    timestamp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in timestamp.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = soft_gray

    divider = doc.add_paragraph("────────────────────────────────────────────")
    divider.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in divider.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = soft_gray
    doc.add_paragraph("")

    # -------------------------------
    # Formatting helpers
    # -------------------------------
    def add_section_heading(text):
        p = doc.add_paragraph(text.strip())
        for run in p.runs:
            run.font.name = base_font
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = heading_color
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_subheading(text):
        p = doc.add_paragraph(text.strip())
        for run in p.runs:
            run.font.name = base_font
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = heading_color
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text.strip(), style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = base_font
            run.font.size = Pt(11)
            run.font.color.rgb = body_color
        return p

    def add_text_block(text):
        p = doc.add_paragraph(text.strip())
        for run in p.runs:
            run.font.name = base_font
            run.font.size = Pt(11)
            run.font.color.rgb = body_color
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_quote(text):
        p = doc.add_paragraph(f"“{text.strip()}”")
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(90, 90, 90)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_divider():
        p = doc.add_paragraph("────────────────────────────────────────────")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = soft_gray
        return p

    def add_code_block(code_text):
        """Adds a monospaced gray-background code block (Word-safe)."""
        style_name = "CodeBlock"
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = "Consolas"
            style.font.size = Pt(10)
            style.paragraph_format.left_indent = Inches(0.4)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)

        p = doc.add_paragraph(code_text, style=style_name)

        # ✅ Proper way to add shading (background fill)
        shading_elm = parse_xml(r'<w:shd {} w:fill="EDEDED"/>'.format(nsdecls('w')))
        p._element.get_or_add_pPr().append(shading_elm)

        return p

    # -------------------------------
    # ENM Parsing (main logic)
    # -------------------------------
    in_code_block = False
    code_buffer = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # Handle start/end of code block
        if line.startswith("#CODE"):
            in_code_block = True
            code_buffer = []
            continue
        if line.startswith("#ENDCODE"):
            in_code_block = False
            if code_buffer:
                add_code_block("\n".join(code_buffer))
                code_buffer = []
            continue

        # Inside a code block → raw text accumulation
        if in_code_block:
            clean = re.sub(r"^#POINT\s*", "", line)
            code_buffer.append(clean)
            continue

        # Normal ENM tags
        if line.startswith("#SECTION"):
            add_section_heading(line.replace("#SECTION", "").strip())
        elif line.startswith("#SUBSECTION"):
            add_subheading(line.replace("#SUBSECTION", "").strip())
        elif line.startswith("#POINT"):
            add_bullet(line.replace("#POINT", "").strip())
        elif line.startswith("#NOTE"):
            add_text_block(line.replace("#NOTE", "").strip())
        elif line.startswith("#QUOTE"):
            add_quote(line.replace("#QUOTE", "").strip())
        elif line.startswith("#DIVIDER"):
            add_divider()
        else:
            # Fallback for any text without tags
            add_text_block(line)

    # -------------------------------
    # Footer
    # -------------------------------
    footer = section.footer.paragraphs[0]
    footer.text = "Generated by EasyNote AI • Everyday Intelligence"
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.name = base_font
        run.italic = True
        run.font.color.rgb = soft_gray

    # -------------------------------
    # Save document
    # -------------------------------
    temp_dir = tempfile.gettempdir()
    safe_title = title.replace(" ", "_")
    file_path = os.path.join(temp_dir, f"{safe_title}.docx")
    doc.save(file_path)
    return file_path
